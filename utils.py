import pandas as pd
from typing import Dict
import datetime

from docxtpl import DocxTemplate
import jinja2
from docx import Document


import warnings

warnings.filterwarnings("ignore")


def _read_data(data_dir: str = "Activities.csv") -> pd.DataFrame:
    df = pd.read_csv(data_dir)
    return df


def _preprocess(df: pd.DataFrame, start_day: str, number_of_days: int) -> pd.DataFrame:
    df = _create_date_and_time_columns(df)
    df = _filter_by_date(df, start_day, number_of_days)
    df["am_pm"] = df["Time_of_Day"].apply(_morning_or_afternoon)
    return df


def _filter_by_date(df: pd.DataFrame, start_day: str, number_of_days: int) -> pd.DataFrame:
    start_datetime = pd.to_datetime(start_day)  # Convert start_day to pandas Timestamp
    date_filter = start_datetime - pd.Timedelta(days=number_of_days)
    return df[(df["Date"] >= date_filter.date()) & (df["Date"] <= start_datetime.date())]


def _create_date_and_time_columns(df: pd.DataFrame) -> pd.DataFrame:
    df["Date"] = pd.to_datetime(df["Date"])
    df["Time_of_Day"] = df["Date"].dt.time
    df["Date"] = df["Date"].dt.date
    return df


def _morning_or_afternoon(time: pd.Timestamp) -> str:
    if time.hour < 12:
        return "morning"
    else:
        return "afternoon"


def prepare_data(data_dir: str, number_of_days: int, start_day: str) -> pd.DataFrame:
    df = _read_data(data_dir)
    df = _preprocess(df, start_day, number_of_days)
    return df

def replace_text(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if old_text in run.text:
            print(f"Replacing '{old_text}' with '{new_text}'")  # Debug statement
            run.text = run.text.replace(old_text, new_text)

def fill_template(data, template_path, output_path):
    document = Document(template_path)

    days_of_week = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY", "SUNDAY"]
    start_date = datetime.date(2024, 6, 3)
    total_distance = 0.0

    for i, day in enumerate(days_of_week):
        current_date = start_date + datetime.timedelta(days=i)
        date_str = current_date.strftime("%Y-%m-%d")

        for paragraph in document.paragraphs:
            replace_text(paragraph, f"[DATE_{day}]", date_str)

        morning_data = data.get(current_date, {}).get('morning', [])
        afternoon_data = data.get(current_date, {}).get('afternoon', [])

        morning_time = ", ".join([item['Time'] for item in morning_data])
        morning_dist = ", ".join([str(item['Distance']) for item in morning_data])
        morning_pace = ", ".join([item['Pace'] for item in morning_data])

        afternoon_time = ", ".join([item['Time'] for item in afternoon_data])
        afternoon_dist = ", ".join([str(item['Distance']) for item in afternoon_data])
        afternoon_pace = ", ".join([item['Pace'] for item in afternoon_data])

        for paragraph in document.paragraphs:
            replace_text(paragraph, f"[TIME_{day}_MORN]", morning_time)
            replace_text(paragraph, f"[DIST_{day}_MORN]", morning_dist)
            replace_text(paragraph, f"[PACE_{day}_MORN]", morning_pace)

            replace_text(paragraph, f"[TIME_{day}_AFTER]", afternoon_time)
            replace_text(paragraph, f"[DIST_{day}_AFTER]", afternoon_dist)
            replace_text(paragraph, f"[PACE_{day}_AFTER]", afternoon_pace)

        total_distance += sum(item['Distance'] for item in morning_data + afternoon_data)

    for paragraph in document.paragraphs:
        replace_text(paragraph, "[WEEKLY_DISTANCE]", str(total_distance))

    document.save(output_path)


# Function to prepare context for the template
def prepare_context(data):
    context = {}
    days_of_week = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY", "SUNDAY"]
    start_date = datetime.date(2024, 6, 3)
    total_distance = 0.0

    for i, day in enumerate(days_of_week):
        current_date = start_date + datetime.timedelta(days=i)
        date_str = current_date.strftime("%Y-%m-%d")

        morning_data = data.get(current_date, {}).get('morning', [])
        afternoon_data = data.get(current_date, {}).get('afternoon', [])

        morning_time = ", ".join([item['Time'] for item in morning_data]) if morning_data else ""
        morning_dist = ", ".join([str(item['Distance']) for item in morning_data]) if morning_data else ""
        morning_pace = ", ".join([item['Pace'] for item in morning_data]) if morning_data else ""

        afternoon_time = ", ".join([item['Time'] for item in afternoon_data]) if afternoon_data else ""
        afternoon_dist = ", ".join([str(item['Distance']) for item in afternoon_data]) if afternoon_data else ""
        afternoon_pace = ", ".join([item['Pace'] for item in afternoon_data]) if afternoon_data else ""

        context[f"DATE_{day}"] = date_str
        context[f"TIME_{day}_MORN"] = morning_time
        context[f"DIST_{day}_MORN"] = morning_dist
        context[f"PACE_{day}_MORN"] = morning_pace
        context[f"TIME_{day}_AFTER"] = afternoon_time
        context[f"DIST_{day}_AFTER"] = afternoon_dist
        context[f"PACE_{day}_AFTER"] = afternoon_pace

        print(f"{day}: Date: {date_str}, Morning: {morning_time}, {morning_dist}, {morning_pace}, Afternoon: {afternoon_time}, {afternoon_dist}, {afternoon_pace}")  # Debug statement

        total_distance += sum(item['Distance'] for item in morning_data + afternoon_data)

    context["WEEKLY_DISTANCE"] = str(round(total_distance,2))
    print(f"Weekly Total Distance: {total_distance}")  # Debug statement

    return context
